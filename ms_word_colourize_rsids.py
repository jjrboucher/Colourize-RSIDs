import tkinter as tk
from tkinter import ttk, filedialog, colorchooser, messagebox, scrolledtext
import zipfile
import xml.etree.ElementTree as ET
import re
import os
import tempfile
import shutil
import logging
import atexit
import hashlib
import platform
import socket
import getpass
from datetime import datetime, timezone
from docx import Document

try:
    import win32com.client
except ImportError:
    win32com = None

APP_NAME = 'Word rsidR Colorizer'
APP_VERSION = '1.2'

# ---------------------------------------------------------------------------
# Logger
# ---------------------------------------------------------------------------
logging.basicConfig(level=logging.DEBUG,
                    format='%(asctime)s - %(levelname)s - %(message)s')
logger = logging.getLogger(__name__)

# ---------------------------------------------------------------------------
# Colour palette  (Microsoft Word–inspired)
# ---------------------------------------------------------------------------
C = {
    'ribbon_bg':         '#2B579A',
    'ribbon_btn_hover':  '#1E4480',
    'ribbon_btn_press':  '#16336B',
    'ribbon_fg':         '#FFFFFF',
    'sidebar_bg':        '#EFF3F8',
    'sidebar_border':    '#C7D3E8',
    'sidebar_header_bg': '#2B579A',
    'sidebar_header_fg': '#FFFFFF',
    'row_bg':            '#FFFFFF',
    'row_bg_alt':        '#F5F8FD',
    'row_hover':         '#DCE8F8',
    'row_border':        '#D6E0EF',
    'row_fg':            '#1A1A1A',
    'swatch_empty':      '#E0E0E0',
    'reset_fg':          '#2B579A',
    'preview_bg':        '#F1F1F1',
    'page_bg':           '#FFFFFF',
    'page_shadow':       '#CCCCCC',
    'status_bg':         '#2B579A',
    'status_fg':         '#FFFFFF',
    'app_bg':            '#F3F3F3',
    'divider':           '#C7D3E8',
    'text_muted':        '#5A6A80',
}

# ---------------------------------------------------------------------------
# XML helpers
# ---------------------------------------------------------------------------
def _extract_namespace_map(xml_path):
    ns_map = {}
    try:
        for event, elem in ET.iterparse(xml_path, events=['start-ns']):
            prefix, uri = elem
            ns_map[prefix] = uri
    except Exception as exc:
        logger.warning(f"Could not extract namespace map from {xml_path}: {exc}")
    return ns_map

def _inject_color_into_run_xml(run_xml, hex_color):
    color_tag = f'<w:color w:val="{hex_color}"/>'
    if '<w:rPr>' in run_xml or '<w:rPr ' in run_xml:
        run_xml = re.sub(r'<w:color\b[^/]*/>', '', run_xml)
        run_xml = re.sub(r'(<w:rPr(?:\s[^>]*)?>\s*)',
                         r'\1' + color_tag, run_xml, count=1)
    else:
        run_xml = re.sub(r'(<w:r(?:\s[^>]*)?>)',
                         r'\1<w:rPr>' + color_tag + '</w:rPr>',
                         run_xml, count=1)
    return run_xml

# ---------------------------------------------------------------------------
# Forensic helpers
# ---------------------------------------------------------------------------
CORE_NS = {
    'cp':      'http://schemas.openxmlformats.org/package/2006/metadata/core-properties',
    'dc':      'http://purl.org/dc/elements/1.1/',
    'dcterms': 'http://purl.org/dc/terms/',
}
APP_NS_URI = 'http://schemas.openxmlformats.org/officeDocument/2006/extended-properties'

def _compute_file_hashes(path):
    """Return a dict with sha256, md5 and size for the file at *path*."""
    sha256 = hashlib.sha256()
    md5 = hashlib.md5()
    size = 0
    with open(path, 'rb') as f:
        for chunk in iter(lambda: f.read(65536), b''):
            sha256.update(chunk)
            md5.update(chunk)
            size += len(chunk)
    return {
        'sha256': sha256.hexdigest().upper(),
        'md5':    md5.hexdigest().upper(),
        'size':   size,
    }

def _get_text(parent, xpath, ns):
    """Helper: return the .text of the first matching element, or None."""
    if parent is None:
        return None
    elem = parent.find(xpath, ns)
    if elem is None or elem.text is None:
        return None
    return elem.text.strip() or None

def _extract_core_metadata(extracted_dir):
    """Parse docProps/core.xml; returns a dict. Missing fields are None."""
    path = os.path.join(extracted_dir, 'docProps', 'core.xml')
    meta = {
        'title':            None,
        'subject':          None,
        'creator':          None,
        'keywords':         None,
        'description':      None,
        'last_modified_by': None,
        'revision':         None,
        'created':          None,
        'modified':         None,
        'category':         None,
        'content_status':   None,
    }
    if not os.path.isfile(path):
        return meta
    try:
        root = ET.parse(path).getroot()
        meta['title']            = _get_text(root, 'dc:title',           CORE_NS)
        meta['subject']          = _get_text(root, 'dc:subject',         CORE_NS)
        meta['creator']          = _get_text(root, 'dc:creator',         CORE_NS)
        meta['keywords']         = _get_text(root, 'cp:keywords',        CORE_NS)
        meta['description']      = _get_text(root, 'dc:description',     CORE_NS)
        meta['last_modified_by'] = _get_text(root, 'cp:lastModifiedBy',  CORE_NS)
        meta['revision']         = _get_text(root, 'cp:revision',         CORE_NS)
        meta['created']          = _get_text(root, 'dcterms:created',    CORE_NS)
        meta['modified']         = _get_text(root, 'dcterms:modified',   CORE_NS)
        meta['category']         = _get_text(root, 'cp:category',        CORE_NS)
        meta['content_status']   = _get_text(root, 'cp:contentStatus',   CORE_NS)
    except Exception as exc:
        logger.warning(f'Failed to parse core.xml: {exc}')
    return meta

def _extract_app_metadata(extracted_dir):
    """Parse docProps/app.xml; returns a dict. Missing fields are None."""
    path = os.path.join(extracted_dir, 'docProps', 'app.xml')
    fields = ['Application', 'AppVersion', 'Company', 'Manager',
              'Template', 'TotalTime', 'Pages', 'Words', 'Characters',
              'CharactersWithSpaces', 'Lines', 'Paragraphs', 'DocSecurity']
    meta = {f: None for f in fields}
    if not os.path.isfile(path):
        return meta
    try:
        root = ET.parse(path).getroot()
        ns = {'a': APP_NS_URI}
        for f in fields:
            meta[f] = _get_text(root, f'a:{f}', ns)
    except Exception as exc:
        logger.warning(f'Failed to parse app.xml: {exc}')
    return meta

def _format_value(v):
    """Render a metadata value for the log; show '(not set)' if blank."""
    if v is None or v == '':
        return '(not set)'
    return v

def _now_utc_iso():
    return datetime.now(timezone.utc).strftime('%Y-%m-%d %H:%M:%S UTC')

def _now_utc_filename_stamp():
    return datetime.now(timezone.utc).strftime('%Y%m%d-%H%M%S')

def _hex_to_rgb_triple(hex_color):
    """Convert '#RRGGBB' to '(R, G, B)' string."""
    h = hex_color.lstrip('#')
    try:
        r, g, b = int(h[0:2], 16), int(h[2:4], 16), int(h[4:6], 16)
        return f'({r}, {g}, {b})'
    except Exception:
        return '(?, ?, ?)'

def _section(title):
    """Format a forensic-style section banner."""
    dashes = '-' * max(0, 80 - len(title) - 6)
    return f'----- {title} {dashes}'

# ---------------------------------------------------------------------------
# Custom widgets
# ---------------------------------------------------------------------------
class RibbonButton(tk.Button):
    def __init__(self, parent, **kw):
        kw.setdefault('bg', C['ribbon_bg'])
        kw.setdefault('fg', C['ribbon_fg'])
        kw.setdefault('activebackground', C['ribbon_btn_hover'])
        kw.setdefault('activeforeground', C['ribbon_fg'])
        kw.setdefault('relief', tk.FLAT)
        kw.setdefault('bd', 0)
        kw.setdefault('padx', 16)
        kw.setdefault('pady', 6)
        kw.setdefault('cursor', 'hand2')
        kw.setdefault('font', ('Segoe UI', 9))
        super().__init__(parent, **kw)
        self.bind('<Enter>', lambda e: self.config(bg=C['ribbon_btn_hover']))
        self.bind('<Leave>', lambda e: self.config(bg=C['ribbon_bg']))
        self.bind('<ButtonPress-1>', lambda e: self.config(bg=C['ribbon_btn_press']))
        self.bind('<ButtonRelease-1>', lambda e: self.config(bg=C['ribbon_btn_hover']))

class SidebarButton(tk.Button):
    def __init__(self, parent, **kw):
        kw.setdefault('bg', C['ribbon_bg'])
        kw.setdefault('fg', C['ribbon_fg'])
        kw.setdefault('activebackground', C['ribbon_btn_hover'])
        kw.setdefault('activeforeground', C['ribbon_fg'])
        kw.setdefault('relief', tk.FLAT)
        kw.setdefault('bd', 0)
        kw.setdefault('padx', 10)
        kw.setdefault('pady', 5)
        kw.setdefault('cursor', 'hand2')
        kw.setdefault('font', ('Segoe UI', 8))
        super().__init__(parent, **kw)
        self.bind('<Enter>', lambda e: self.config(bg=C['ribbon_btn_hover']))
        self.bind('<Leave>', lambda e: self.config(bg=C['ribbon_bg']))

class RsidRow(tk.Frame):
    def __init__(self, parent, rsid, count, on_pick, on_reset, row_index, **kw):
        bg = C['row_bg'] if row_index % 2 == 0 else C['row_bg_alt']
        super().__init__(parent, bg=bg, **kw)
        self._bg = bg
        self.rsid = rsid

        self.swatch = tk.Canvas(self, width=22, height=22,
                                bg=bg, highlightthickness=1,
                                highlightbackground=C['row_border'],
                                cursor='hand2')
        self.swatch.pack(side=tk.LEFT, padx=(8, 6), pady=6)
        self._draw_swatch(C['swatch_empty'])
        self.swatch.bind('<Button-1>', lambda e: on_pick(rsid))

        self.lbl_frame = tk.Frame(self, bg=bg)
        self.lbl_frame.pack(side=tk.LEFT, fill=tk.X, expand=True, pady=6)

        self.rsid_lbl = tk.Label(self.lbl_frame, text=rsid,
                                 font=('Consolas', 9, 'bold'),
                                 bg=bg, fg=C['row_fg'], cursor='hand2')
        self.rsid_lbl.pack(anchor=tk.W)
        self.rsid_lbl.bind('<Button-1>', lambda e: on_pick(rsid))

        self.count_lbl = tk.Label(self.lbl_frame,
                                  text=f'{count} occurrence{"s" if count != 1 else ""}',
                                  font=('Segoe UI', 7),
                                  bg=bg, fg=C['text_muted'])
        self.count_lbl.pack(anchor=tk.W)

        self.reset_lbl = tk.Label(self, text='Reset',
                                  font=('Segoe UI', 8, 'underline'),
                                  bg=bg, fg=C['reset_fg'], cursor='hand2')
        self.reset_lbl.pack(side=tk.RIGHT, padx=(4, 10))
        self.reset_lbl.bind('<Button-1>', lambda e: on_reset(rsid))

        for widget in (self, self.lbl_frame, self.rsid_lbl,
                       self.count_lbl, self.swatch, self.reset_lbl):
            widget.bind('<Enter>', self._on_enter)
            widget.bind('<Leave>', self._on_leave)

    def _draw_swatch(self, color):
        self.swatch.delete('all')
        self.swatch.create_rectangle(2, 2, 20, 20, fill=color,
                                     outline=C['row_border'])

    def set_color(self, hex_color):
        self._draw_swatch(hex_color)
        self.swatch.config(highlightbackground='#888888')

    def clear_color(self):
        self._draw_swatch(C['swatch_empty'])
        self.swatch.config(highlightbackground=C['row_border'])

    def _on_enter(self, _e):
        self._set_row_bg(C['row_hover'])

    def _on_leave(self, _e):
        # <Leave> also fires when the pointer moves from the row onto one of
        # its own child widgets. Only revert to the base colour if the pointer
        # has genuinely left this row, otherwise the background flickers
        # back to white as the mouse moves across the row.
        x, y = self.winfo_pointerxy()
        widget = self.winfo_containing(x, y)
        while widget is not None:
            if widget is self:
                return
            widget = getattr(widget, 'master', None)
        self._set_row_bg(self._bg)

    def _set_row_bg(self, color):
        # Recolour every part of the row, including lbl_frame (the container
        # behind the rsid/count labels) so the hover fill is solid with no
        # white showing through.
        for w in (self, self.lbl_frame, self.swatch,
                  self.rsid_lbl, self.count_lbl, self.reset_lbl):
            try:
                w.config(bg=color)
            except Exception:
                pass

# ---------------------------------------------------------------------------
# Main WordRsidColorizer Application
# ---------------------------------------------------------------------------
class WordRsidColorizer:
    def __init__(self, root):
        self.root = root
        self.root.title(APP_NAME)
        self.root.geometry('1280x820')
        self.root.minsize(900, 600)
        self.root.configure(bg=C['app_bg'])

        self.current_file = None
        self.last_save_path = None
        self.rsid_colors = {}
        self.rsid_rows = {}
        self.rsid_counts = {}            
        self.document_xml = None
        self.document_tree = None
        self.document_ns_map = {}
        self.temp_dir = None
        self._cached_page_map = {}

        # Forensic state
        self.source_hashes = None        
        self.source_metadata_core = {}
        self.source_metadata_app = {}

        self.namespaces = {
            'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'
        }
        ET.register_namespace('w', self.namespaces['w'])

        self._build_ui()
        atexit.register(self._cleanup_temp)

    def _build_menu(self):
        """Build the application window top menubar hierarchy."""
        self.menubar = tk.Menu(self.root)
        self.report_menu = tk.Menu(self.menubar, tearoff=0)
        self.report_menu.add_command(label='Generate RSID Report…', command=self.generate_report)
        self.menubar.add_cascade(label='Reports', menu=self.report_menu)
        self.root.config(menu=self.menubar)


    def _build_paragraph_page_map(self):
        """
        Returns {paragraph_number: page_number}. Requires Microsoft Word and pywin32.
        The mapping is generated once per report and cached to minimize COM calls.
        """
        if hasattr(self, "_cached_page_map") and self._cached_page_map:
            return self._cached_page_map

        page_map = {}
        if win32com is None or not self.current_file:
            return page_map

        word = None
        doc = None
        try:
            word = win32com.client.DispatchEx('Word.Application')
            word.Visible = False
            word.ScreenUpdating = False
            word.DisplayAlerts = False

            doc = word.Documents.Open(
                os.path.abspath(self.current_file),
                ReadOnly=True,
                AddToRecentFiles=False,
                Visible=False
            )

            wdActiveEndAdjustedPageNumber = 1
            para_count = doc.Paragraphs.Count

            for idx in range(1, para_count + 1):
                page_map[idx] = doc.Paragraphs(idx).Range.Information(
                    wdActiveEndAdjustedPageNumber
                )

            self._cached_page_map = page_map

        except Exception as exc:
            logger.warning(f'Unable to determine page numbers: {exc}')
        finally:
            if doc:
                doc.Close(False)
            if word:
                word.Quit()

        return page_map

        word = None
        doc = None
        try:
            word = win32com.client.Dispatch('Word.Application')
            word.Visible = False
            doc = word.Documents.Open(os.path.abspath(self.current_file), ReadOnly=True)

            wdActiveEndAdjustedPageNumber = 1
            for idx in range(1, doc.Paragraphs.Count + 1):
                rng = doc.Paragraphs(idx).Range
                page_map[idx] = rng.Information(wdActiveEndAdjustedPageNumber)
        except Exception as exc:
            logger.warning(f'Unable to determine page numbers: {exc}')
        finally:
            if doc:
                doc.Close(False)
            if word:
                word.Quit()

        return page_map

    def generate_report(self):
        """Forensic extraction reporting out RSID blocks indexed via Paragraph ordering."""
        if not self.current_file or not self.document_tree:
            messagebox.showwarning('Warning', 'No document loaded.')
            return
        
        save_path = filedialog.asksaveasfilename(
            title='Save RSID Report',
            defaultextension='.docx',
            filetypes=[('Word Documents', '*.docx'), ('All Files', '*.*')])
        if not save_path:
            return

        try:
            doc = Document()
            doc.add_heading(f'RSID Forensic Mapping Report', level=0)
            doc.add_paragraph(f'Source Document File: {os.path.basename(self.current_file)}')
                        
            # Extract and inject the MD5 hash into the report body
            md5_str = self.source_hashes['md5'] if self.source_hashes else 'Unknown'
            doc.add_paragraph(f'Source Document MD5: {md5_str}')
            
            doc.add_paragraph(f'Generated on Reference Time: {datetime.now(timezone.utc).strftime("%Y-%m-%d %H:%M:%S UTC")}')
            
            page_map = self._build_paragraph_page_map()
            active_rsids = sorted([r for r, c in self.rsid_counts.items() if c > 0])
            w_ns = self.namespaces['w']
            root = self.document_tree.getroot()
            
            for rsid in active_rsids:
                paragraphs_found = []
                para_num = 0
                
                for para in root.findall('.//w:p', self.namespaces):
                    para_num += 1
                    para_rsid = para.get(f'{{{w_ns}}}rsidR')
                    para_rsid_p = para.get(f'{{{w_ns}}}rsidP')
                    
                    match_found = False
                    para_text_parts = []
                    
                    for run in para.findall('.//w:r', self.namespaces):
                        run_rsid = run.get(f'{{{w_ns}}}rsidR')
                        effective_rsid = run_rsid or para_rsid or para_rsid_p
                        
                        if effective_rsid == rsid:
                            match_found = True
                            parts = []
                            for t in run.findall('./w:t', self.namespaces):
                                if t.text:
                                    parts.append(t.text)
                            for _ in run.findall('./w:tab', self.namespaces):
                                parts.append('\t')
                            for _ in run.findall('./w:br', self.namespaces):
                                parts.append('\n')
                            if parts:
                                para_text_parts.append(''.join(parts))
                                
                    if not match_found and (para_rsid == rsid or para_rsid_p == rsid):
                        full_parts = []
                        for run in para.findall('.//w:r', self.namespaces):
                            for t in run.findall('./w:t', self.namespaces):
                                if t.text:
                                    full_parts.append(t.text)
                        if full_parts:
                            para_text_parts.append(''.join(full_parts))
                            match_found = True
                            
                    if match_found and para_text_parts:
                        combined_text = ''.join(para_text_parts).strip()
                        if combined_text:
                            paragraphs_found.append((page_map.get(para_num, '?'), para_num, combined_text))
                            
                if paragraphs_found:
                    doc.add_heading(f'RSID Value Reference: {rsid}', level=1)
                    table = doc.add_table(rows=1, cols=3)
                    table.style = 'Table Grid'
                    hdr_cells = table.rows[0].cells
                    hdr_cells[0].text = 'Page #'
                    hdr_cells[1].text = 'Paragraph #'
                    hdr_cells[2].text = 'Extracted Text Segments'
                    
                    for page_num, p_num, text in paragraphs_found:
                        row_cells = table.add_row().cells
                        row_cells[0].text = str(page_num)
                        row_cells[1].text = str(p_num)
                        row_cells[2].text = text
            
            doc.save(save_path)
            messagebox.showinfo('Success', f'Forensic RSID Mapping Report saved successfully to:\n{save_path}')
        except Exception as exc:
            logger.exception('Failed to generate report')
            messagebox.showerror('Report Error', f'Failed to generate report framework:\n{exc}')

    def _build_ui(self):
        ribbon = tk.Frame(self.root, bg=C['ribbon_bg'], height=48)
        ribbon.pack(fill=tk.X, side=tk.TOP)
        ribbon.pack_propagate(False)

        tk.Label(ribbon, text='W  rsidR Colorizer',
                 font=('Segoe UI', 12, 'bold'),
                 bg=C['ribbon_bg'], fg=C['ribbon_fg'],
                 padx=16).pack(side=tk.LEFT, pady=8)

        tk.Frame(ribbon, bg='#4A72B0', width=1).pack(side=tk.LEFT, fill=tk.Y, pady=8)

        RibbonButton(ribbon, text='Open Document', command=self.select_file).pack(side=tk.LEFT, padx=(12, 2), pady=8)
        RibbonButton(ribbon, text='Save', command=self.save_document).pack(side=tk.LEFT, padx=2, pady=8)
        RibbonButton(ribbon, text='Save As…', command=self.save_document_as).pack(side=tk.LEFT, padx=2, pady=8)
        RibbonButton(ribbon, text='Generate Report', command=self.generate_report).pack(side=tk.LEFT, padx=(12,2), pady=8)

        status_bar = tk.Frame(self.root, bg=C['status_bg'], height=26)
        status_bar.pack(fill=tk.X, side=tk.BOTTOM)
        status_bar.pack_propagate(False)

        self.status_var = tk.StringVar(value='No document loaded')
        tk.Label(status_bar, textvariable=self.status_var,
                 font=('Segoe UI', 8), bg=C['status_bg'],
                 fg=C['status_fg'], padx=12, anchor=tk.W).pack(fill=tk.X, expand=True, pady=4)

        self.file_strip = tk.Frame(self.root, bg='#1E4480', height=28)
        self.file_strip.pack(fill=tk.X)
        self.file_strip.pack_propagate(False)

        self.file_label = tk.Label(self.file_strip, text='', font=('Segoe UI', 8),
                                   bg='#1E4480', fg='#BED0F0', padx=14, anchor=tk.W)
        self.file_label.pack(fill=tk.X, expand=True, pady=5)

        content = tk.Frame(self.root, bg=C['app_bg'])
        content.pack(fill=tk.BOTH, expand=True)

        self._build_sidebar(content)
        tk.Frame(content, bg=C['divider'], width=1).pack(side=tk.LEFT, fill=tk.Y)
        self._build_preview(content)

    def _build_sidebar(self, parent):
        sidebar = tk.Frame(parent, bg=C['sidebar_bg'], width=290)
        sidebar.pack(side=tk.LEFT, fill=tk.Y)
        sidebar.pack_propagate(False)

        hdr = tk.Frame(sidebar, bg=C['sidebar_header_bg'], height=36)
        hdr.pack(fill=tk.X)
        hdr.pack_propagate(False)
        tk.Label(hdr, text='rsidR VALUES', font=('Segoe UI', 8, 'bold'),
                 bg=C['sidebar_header_bg'], fg=C['sidebar_header_fg'], padx=12).pack(side=tk.LEFT, pady=8)

        SidebarButton(hdr, text='Clear All', command=self.clear_all_colors, padx=8, pady=3,
                      font=('Segoe UI', 8)).pack(side=tk.RIGHT, padx=8, pady=6)

        scroll_wrapper = tk.Frame(sidebar, bg=C['sidebar_bg'])
        scroll_wrapper.pack(fill=tk.BOTH, expand=True)

        self.canvas = tk.Canvas(scroll_wrapper, bg=C['sidebar_bg'], highlightthickness=0)
        vsb = ttk.Scrollbar(scroll_wrapper, orient='vertical', command=self.canvas.yview)
        self.scrollable_frame = tk.Frame(self.canvas, bg=C['sidebar_bg'])

        self.scrollable_frame.bind('<Configure>', lambda e: self.canvas.configure(scrollregion=self.canvas.bbox('all')))
        self.canvas_window = self.canvas.create_window((0, 0), window=self.scrollable_frame, anchor='nw')
        self.canvas.configure(yscrollcommand=vsb.set)

        self.canvas.pack(side='left', fill='both', expand=True)
        vsb.pack(side='right', fill='y')

        self.canvas.bind('<MouseWheel>', self._on_mousewheel)
        self.canvas.bind('<Button-4>', self._on_mousewheel)
        self.canvas.bind('<Button-5>', self._on_mousewheel)
        self.canvas.bind('<Configure>', self._on_canvas_configure)

        self.sidebar_placeholder = tk.Label(
            self.scrollable_frame, text='Open a document to\nview rsidR values.',
            font=('Segoe UI', 9), bg=C['sidebar_bg'], fg=C['text_muted'], justify=tk.CENTER)
        self.sidebar_placeholder.pack(pady=40)

    def _build_preview(self, parent):
        preview_outer = tk.Frame(parent, bg=C['preview_bg'])
        preview_outer.pack(side=tk.RIGHT, fill=tk.BOTH, expand=True)

        prev_hdr = tk.Frame(preview_outer, bg='#E0E0E0', height=32)
        prev_hdr.pack(fill=tk.X)
        prev_hdr.pack_propagate(False)
        tk.Label(prev_hdr, text='DOCUMENT PREVIEW', font=('Segoe UI', 8, 'bold'),
                 bg='#E0E0E0', fg='#555555', padx=14).pack(side=tk.LEFT, pady=8)

        page_frame = tk.Frame(preview_outer, bg=C['preview_bg'])
        page_frame.pack(fill=tk.BOTH, expand=True, padx=20, pady=16)

        shadow = tk.Frame(page_frame, bg=C['page_shadow'])
        shadow.pack(fill=tk.BOTH, expand=True)

        page = tk.Frame(shadow, bg=C['page_bg'])
        page.pack(fill=tk.BOTH, expand=True, padx=2, pady=2)

        self.preview_text = scrolledtext.ScrolledText(
            page, wrap=tk.WORD, state=tk.DISABLED, font=('Cambria', 11), bg=C['page_bg'], fg='#1A1A1A',
            relief=tk.FLAT, bd=0, padx=32, pady=24, selectbackground='#BDD7EE')
        self.preview_text.pack(fill=tk.BOTH, expand=True)

    def _on_mousewheel(self, event):
        if event.delta:
            self.canvas.yview_scroll(int(-1 * (event.delta / 120)), 'units')
        else:
            if event.num == 4:
                self.canvas.yview_scroll(-1, 'units')
            elif event.num == 5:
                self.canvas.yview_scroll(1, 'units')

    def _on_canvas_configure(self, event):
        self.canvas.itemconfig(self.canvas_window, width=event.width)

    def _bind_mousewheel_recursive(self, widget):
        """Recursively binds scroll events up to the canvas component framework."""
        widget.bind('<MouseWheel>', self._on_mousewheel)
        widget.bind('<Button-4>', self._on_mousewheel)
        widget.bind('<Button-5>', self._on_mousewheel)
        for child in widget.winfo_children():
            self._bind_mousewheel_recursive(child)

    def select_file(self):
        path = filedialog.askopenfilename(
            title='Select Word Document',
            filetypes=[('Word Documents', '*.docx'), ('All Files', '*.*')])
        if path:
            self.current_file = os.path.abspath(path)
            self.last_save_path = None
            self.rsid_colors.clear()
            filename = os.path.basename(self.current_file)
            self.file_label.config(text=f'  {filename}')
            self.root.title(f'{APP_NAME}  –  {filename}')
            self.load_document()

    def load_document(self):
        try:
            if self.temp_dir:
                shutil.rmtree(self.temp_dir, ignore_errors=True)
            self.temp_dir = tempfile.mkdtemp()

            self.source_hashes = _compute_file_hashes(self.current_file)

            with zipfile.ZipFile(self.current_file, 'r') as z:
                z.extractall(self.temp_dir)

            self.file_label.config(text=f'  {os.path.basename(self.current_file)}    |    MD5: {self.source_hashes["md5"]}')

            self.source_metadata_core = _extract_core_metadata(self.temp_dir)
            self.source_metadata_app  = _extract_app_metadata(self.temp_dir)

            settings_path = os.path.join(self.temp_dir, 'word', 'settings.xml')
            rsid_values = self.extract_rsid_values(settings_path)
            doc_path = os.path.join(self.temp_dir, 'word', 'document.xml')
            self.load_document_xml(doc_path)

            self.rsid_counts = self.count_rsid_usage(rsid_values)
            self.create_rsid_rows(rsid_values)
            self.update_preview()
            self._update_status()

            log_path = self._write_examination_log()
            if log_path:
                self.status_var.set(
                    f'  {self.status_var.get().strip()}   •   '
                    f'Examination log: {os.path.basename(log_path)}')
        except Exception as exc:
            logger.exception('Failed to load document')
            messagebox.showerror('Error', f'Failed to load document:\n{exc}')

    def save_document(self):
        if not self.current_file or not self.document_xml:
            messagebox.showwarning('Warning', 'No document loaded.')
            return
        if self.last_save_path:
            self._write_document(self.last_save_path)
        else:
            self.save_document_as()

    def save_document_as(self):
        if not self.current_file or not self.document_xml:
            messagebox.showwarning('Warning', 'No document loaded.')
            return
        save_path = filedialog.asksaveasfilename(
            title='Save Coloured Document', defaultextension='.docx',
            filetypes=[('Word Documents', '*.docx'), ('All Files', '*.*')])
        if save_path:
            self._write_document(save_path)

    def _write_document(self, save_path):
        try:
            modified_bytes = self.build_modified_document_xml()
            with zipfile.ZipFile(self.current_file, 'r') as zip_in, \
                 zipfile.ZipFile(save_path, 'w') as zip_out:
                for item in zip_in.infolist():
                    if item.filename == 'word/document.xml':
                        info = zipfile.ZipInfo(filename=item.filename, date_time=item.date_time)
                        zip_out.writestr(info, modified_bytes, compress_type=item.compress_type)
                    else:
                        zip_out.writestr(item, zip_in.read(item.filename))

            self.last_save_path = os.path.abspath(save_path)
            logger.info(f'Document saved to {self.last_save_path}')

            output_hashes = _compute_file_hashes(self.last_save_path)
            log_path = self._write_colourization_log(self.last_save_path, output_hashes)
            log_note = f'\n\nForensic log: {log_path}' if log_path else ''
            messagebox.showinfo('Saved', f'Document saved to:\n{self.last_save_path}{log_note}')
        except Exception as exc:
            logger.exception('Failed to save document')
            messagebox.showerror('Save Error', f'Failed to save document:\n{exc}')

    def extract_rsid_values(self, settings_path):
        rsid_values = set()
        try:
            root = ET.parse(settings_path).getroot()
            rsids_elem = root.find('.//w:rsids', self.namespaces)
            if rsids_elem is not None:
                rsid_root = rsids_elem.find('./w:rsidRoot', self.namespaces)
                if rsid_root is not None:
                    val = rsid_root.get(f'{{{self.namespaces["w"]}}}val')
                    if val:
                        rsid_values.add(val)
                for tag in ('w:rsid', 'w:rsidR', 'w:rsidRPr', 'w:rsidDel', 'w:rsidP'):
                    for elem in rsids_elem.findall(f'./{tag}', self.namespaces):
                        val = elem.get(f'{{{self.namespaces["w"]}}}val')
                        if val:
                            rsid_values.add(val)
            else:
                logger.warning('No <w:rsids> element found in settings.xml')
            logger.info(f'Extracted {len(rsid_values)} rsid values')
        except Exception as exc:
            logger.exception('Error extracting rsid values')
        return sorted(rsid_values)

    def load_document_xml(self, doc_path):
        try:
            self.document_ns_map = _extract_namespace_map(doc_path)
            for prefix, uri in self.document_ns_map.items():
                ET.register_namespace(prefix, uri)
            with open(doc_path, 'r', encoding='utf-8') as f:
                self.document_xml = f.read()
            self.document_tree = ET.parse(doc_path)
        except Exception as exc:
            logger.exception('Error loading document.xml')
            raise

    def count_rsid_usage(self, rsid_values):
        counts = {r: 0 for r in rsid_values}
        try:
            if not self.document_tree:
                return counts
            root = self.document_tree.getroot()
            w = self.namespaces['w']
            for elem in root.iter():
                for attr in (f'{{{w}}}rsidR', f'{{{w}}}rsidP', f'{{{w}}}rsidRPr', f'{{{w}}}rsidDel'):
                    val = elem.get(attr)
                    if val and val in counts:
                        counts[val] += 1
        except Exception as exc:
            logger.exception('Error counting rsid usage')
        return counts

    def create_rsid_rows(self, rsid_values):
        for w in self.scrollable_frame.winfo_children():
            w.destroy()
        self.rsid_rows.clear()
        self.sidebar_placeholder = None

        counts = self.rsid_counts
        active = [(r, c) for r, c in counts.items() if c > 0]

        if not active:
            tk.Label(self.scrollable_frame, text='No active rsidR values found.',
                     font=('Segoe UI', 9), bg=C['sidebar_bg'], fg=C['text_muted']).pack(pady=20)
            return

        active.sort(key=lambda x: x[1], reverse=True)
        for idx, (rsid, count) in enumerate(active):
            row = RsidRow(self.scrollable_frame, rsid, count,
                          on_pick=self.select_color, on_reset=self.reset_color, row_index=idx)
            row.pack(fill=tk.X)
            tk.Frame(self.scrollable_frame, bg=C['row_border'], height=1).pack(fill=tk.X)
            self.rsid_rows[rsid] = row
            self._bind_mousewheel_recursive(row)

        self.scrollable_frame.update_idletasks()
        self.canvas.configure(scrollregion=self.canvas.bbox('all'))

    def select_color(self, rsid):
        color = colorchooser.askcolor(title=f'Choose colour for rsid {rsid}')
        if color[1]:
            self.rsid_colors[rsid] = color[1]
            if rsid in self.rsid_rows:
                self.rsid_rows[rsid].set_color(color[1])
            self.update_preview()
            self._update_status()

    def reset_color(self, rsid):
        self.rsid_colors.pop(rsid, None)
        if rsid in self.rsid_rows:
            self.rsid_rows[rsid].clear_color()
        self.update_preview()
        self._update_status()

    def clear_all_colors(self):
        for rsid in list(self.rsid_colors):
            self.rsid_colors.pop(rsid)
            if rsid in self.rsid_rows:
                self.rsid_rows[rsid].clear_color()
        self.update_preview()
        self._update_status()

    def update_preview(self):
        if not self.document_tree:
            return
        try:
            text_content = self.extract_text_with_rsid()
            self.preview_text.config(state=tk.NORMAL)
            self.preview_text.delete(1.0, tk.END)

            for rsid, color in self.rsid_colors.items():
                self.preview_text.tag_configure(f'rsid_{rsid}', foreground=color, font=('Cambria', 11, 'bold'))
            self.preview_text.tag_configure('default', font=('Cambria', 11), foreground='#1A1A1A')

            for text, rsid in text_content:
                tag = f'rsid_{rsid}' if rsid and rsid in self.rsid_colors else 'default'
                self.preview_text.insert(tk.END, text, tag)
            self.preview_text.config(state=tk.DISABLED)
        except Exception as exc:
            logger.exception('Error updating preview')

    def extract_text_with_rsid(self):
        text_content = []
        try:
            root = self.document_tree.getroot()
            w = self.namespaces['w']
            for para in root.findall('.//w:p', self.namespaces):
                para_rsid = para.get(f'{{{w}}}rsidR')
                para_rsid_p = para.get(f'{{{w}}}rsidP')
                para_text = []
                for run in para.findall('.//w:r', self.namespaces):
                    run_rsid = run.get(f'{{{w}}}rsidR')
                    effective = run_rsid or para_rsid or para_rsid_p
                    parts = []
                    for t in run.findall('./w:t', self.namespaces):
                        if t.text:
                            parts.append(t.text)
                    for _ in run.findall('./w:tab', self.namespaces):
                        parts.append('\t')
                    for _ in run.findall('./w:br', self.namespaces):
                        parts.append('\n')
                    if parts:
                        para_text.append((''.join(parts), effective))
                if para_text:
                    text_content.extend(para_text)
                    text_content.append(('\n\n', None))
                elif para_rsid or para_rsid_p:
                    text_content.append(('\n', para_rsid or para_rsid_p))
        except Exception as exc:
            logger.exception('Error extracting text with rsid')
        return text_content

    def build_modified_document_xml(self):
        if not self.document_xml:
            raise RuntimeError('No document XML loaded.')
        if not self.rsid_colors:
            return self.document_xml.encode('utf-8')

        color_map = {rsid: color.lstrip('#').upper() for rsid, color in self.rsid_colors.items()}
        xml = self.document_xml

        def replace_run(m):
            run_xml = m.group(0)
            rm = re.search(r'w:rsidR="([^"]+)"', run_xml)
            if not rm or rm.group(1) not in color_map:
                return run_xml
            return _inject_color_into_run_xml(run_xml, color_map[rm.group(1)])

        xml = re.sub(r'<w:r[ >].*?</w:r>', replace_run, xml, flags=re.DOTALL)

        def replace_para(m):
            para_xml = m.group(0)
            pm = re.search(r'w:rsidP="([^"]+)"', para_xml)
            if not pm or pm.group(1) not in color_map:
                return para_xml
            hex_color = color_map[pm.group(1)]

            def replace_run_in_para(rm2):
                rx = rm2.group(0)
                rr = re.search(r'w:rsidR="([^"]+)"', rx)
                if rr and rr.group(1) in color_map:
                    return rx
                if re.search(r'<w:color\b', rx):
                    return rx
                return _inject_color_into_run_xml(rx, hex_color)

            return re.sub(r'<w:r[ >].*?</w:r>', replace_run_in_para, para_xml, flags=re.DOTALL)

        xml = re.sub(r'<w:p[ >].*?</w:p>', replace_para, xml, flags=re.DOTALL)
        logger.info('document.xml modified successfully')
        return xml.encode('utf-8')

    def _common_log_header(self, log_type):
        try:
            user = getpass.getuser()
        except Exception:
            user = '(unknown)'
        return (
            '=' * 80 + '\n'
            f'{APP_NAME.upper()} — {log_type.upper()} LOG\n'
            + '=' * 80 + '\n'
            f'Log generated:        {_now_utc_iso()}\n'
            f'Tool:                {APP_NAME} v{APP_VERSION}\n'
            f'Python:              {platform.python_version()}\n'
            f'Platform:            {platform.platform()}\n'
            f'Host:                {socket.gethostname()}\n'
            f'OS user:             {user}\n'
            '\n'
        )

    def _format_file_block(self, path, hashes):
        return (
            f'File name:           {os.path.basename(path)}\n'
            f'File path:           {path}\n'
            f'File size:           {hashes["size"]:,} bytes\n'
            f'SHA-256:             {hashes["sha256"]}\n'
            f'MD5:                 {hashes["md5"]}\n'
        )

    def _format_metadata_block(self):
        c = self.source_metadata_core
        a = self.source_metadata_app
        return (
            f'Title:               {_format_value(c.get("title"))}\n'
            f'Subject:             {_format_value(c.get("subject"))}\n'
            f'Author (creator):    {_format_value(c.get("creator"))}\n'
            f'Last modified by:    {_format_value(c.get("last_modified_by"))}\n'
            f'Created:             {_format_value(c.get("created"))}\n'
            f'Last modified:       {_format_value(c.get("modified"))}\n'
            f'Revision number:     {_format_value(c.get("revision"))}\n'
            f'Keywords:            {_format_value(c.get("keywords"))}\n'
            f'Description:         {_format_value(c.get("description"))}\n'
            f'Category:            {_format_value(c.get("category"))}\n'
            f'Content status:      {_format_value(c.get("content_status"))}\n'
            f'Application:         {_format_value(a.get("Application"))}\n'
            f'App version:         {_format_value(a.get("AppVersion"))}\n'
            f'Company:             {_format_value(a.get("Company"))}\n'
            f'Manager:             {_format_value(a.get("Manager"))}\n'
            f'Template:            {_format_value(a.get("Template"))}\n'
            f'Total edit minutes:  {_format_value(a.get("TotalTime"))}\n'
            f'Pages:               {_format_value(a.get("Pages"))}\n'
            f'Words:               {_format_value(a.get("Words"))}\n'
            f'Characters:          {_format_value(a.get("Characters"))}\n'
            f'Document security:   {_format_value(a.get("DocSecurity"))}\n'
        )

    def _format_rsid_table(self):
        active = [(r, c) for r, c in self.rsid_counts.items() if c > 0]
        active.sort(key=lambda x: (-x[1], x[0]))
        total_occ = sum(c for _, c in active)
        out = [f'Unique rsidR values: {len(active)}',
               f'Total occurrences:   {total_occ}',
               '',
               f'{"RSID Value":<14}{"Occurrences":>12}',
               '-' * 26]
        for rsid, count in active:
            out.append(f'{rsid:<14}{count:>12}')
        return '\n'.join(out) + '\n'

    def _write_examination_log(self):
        if not self.current_file or not self.source_hashes:
            return None
        try:
            log_dir = os.path.dirname(self.current_file) or '.'
            base = os.path.splitext(os.path.basename(self.current_file))[0]
            log_path = os.path.join(log_dir, f'{base}_examination_{_now_utc_filename_stamp()}.log')

            content = (
                self._common_log_header('Examination')
                + _section('FILE UNDER EXAMINATION') + '\n'
                + self._format_file_block(self.current_file, self.source_hashes)
                + '\n'
                + _section('DOCUMENT METADATA') + '\n'
                + self._format_metadata_block()
                + '\n'
                + _section('rsidR VALUES DISCOVERED IN DOCUMENT') + '\n'
                + self._format_rsid_table()
                + '\n'
                + '=' * 80 + '\n'
                + 'End of examination log.\n'
            )

            with open(log_path, 'w', encoding='utf-8') as f:
                f.write(content)
            logger.info(f'Examination log written: {log_path}')
            return log_path
        except Exception as exc:
            logger.exception('Failed to write examination log')
            messagebox.showwarning('Log Warning', f'Document opened, but examination log could not be written:\n{exc}')
            return None

    def _write_colourization_log(self, output_path, output_hashes):
        try:
            log_dir = os.path.dirname(output_path) or '.'
            base = os.path.splitext(os.path.basename(output_path))[0]
            log_path = os.path.join(log_dir, f'{base}_colourization_{_now_utc_filename_stamp()}.log')

            if self.rsid_colors:
                rows = [f'{"RSID Value":<14}{"Occurrences":>12}  f{"Colour (hex)":<14}{"RGB":<18}']
                rows.append('-' * 60)
                items = sorted(self.rsid_colors.items(), key=lambda kv: -self.rsid_counts.get(kv[0], 0))
                for rsid, hex_c in items:
                    occ = self.rsid_counts.get(rsid, 0)
                    rows.append(f'{rsid:<14}{occ:>12}  {hex_c.upper():<14}{_hex_to_rgb_triple(hex_c):<18}')
                colour_block = '\n'.join(rows) + '\n'
                total_coloured = sum(self.rsid_counts.get(r, 0) for r in self.rsid_colors)
                summary = (f'rsidR values coloured: {len(self.rsid_colors)} of {len(self.rsid_counts)} discovered\n'
                           f'Affected runs:         {total_coloured}\n\n')
            else:
                colour_block = '(No rsidR values were assigned a colour.)\n'
                summary = ''

            content = (
                self._common_log_header('Colourization')
                + _section('ORIGINAL (SOURCE) FILE') + '\n'
                + self._format_file_block(self.current_file, self.source_hashes)
                + '\n'
                + _section('COLOURIZED (OUTPUT) FILE') + '\n'
                + self._format_file_block(output_path, output_hashes)
                + '\n'
                + _section('ORIGINAL DOCUMENT METADATA') + '\n'
                + self._format_metadata_block()
                + '\n'
                + _section('COLOUR ASSIGNMENTS') + '\n'
                + summary
                + colour_block
                + '\n'
                + '=' * 80 + '\n'
                + 'End of colourization log.\n'
            )

            with open(log_path, 'w', encoding='utf-8') as f:
                f.write(content)
            logger.info(f'Colourization log written: {log_path}')
            return log_path
        except Exception as exc:
            logger.exception('Failed to write colourization log')
            messagebox.showwarning('Log Warning', f'Document saved, but colourization log could not be written:\n{exc}')
            return None

    def _update_status(self):
        total = len(self.rsid_rows)
        coloured = len(self.rsid_colors)
        if total:
            self.status_var.set(f'  {coloured} of {total} rsidR value{"s" if total != 1 else ""} assigned a colour')
        else:
            self.status_var.set('  No document loaded')

    def _cleanup_temp(self):
        if self.temp_dir:
            shutil.rmtree(self.temp_dir, ignore_errors=True)
            self.temp_dir = None
        self._cached_page_map = {}

    def __del__(self):
        self._cleanup_temp()

# ---------------------------------------------------------------------------
def main():
    root = tk.Tk()
    try:
        root.tk.call('tk', 'scaling', 1.25)
    except Exception:
        pass
    app = WordRsidColorizer(root)
    root.mainloop()

if __name__ == '__main__':
    main()